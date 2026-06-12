from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Estilos globales ───────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

GOLD   = RGBColor(0xC9, 0xA9, 0x6E)
DARK   = RGBColor(0x1E, 0x2E, 0x40)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)
DGRAY  = RGBColor(0x37, 0x41, 0x51)

def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def add_border(table):
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'C9A96E')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C9A96E')
    pb.append(bot)
    pPr.append(pb)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK
    hr(doc)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = DGRAY
    return p

def body(doc, text, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10.5)
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F0F4F8')
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size      = Pt(9.5)
    run.font.color.rgb = GRAY
    run.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()

def code_block(doc, text, title=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, '1E2E40')
    p   = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()

def flow_step(doc, num, actor, action):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    w = [Cm(1.2), Cm(3.5), Cm(11)]
    for i, width in enumerate(w):
        table.columns[i].width = width
    c0, c1, c2 = table.row_cells(0)
    set_cell_bg(c0, '1E2E40')
    set_cell_bg(c1, 'F0F4F8')
    for cell, txt in [(c0, str(num)), (c1, actor), (c2, action)]:
        p   = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9.5)
        if cell == c0:
            run.bold = True
            run.font.color.rgb = GOLD
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif cell == c1:
            run.bold = True
            run.font.color.rgb = DARK
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

def uc_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    add_border(table)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(11.5)
    for i, (campo, valor) in enumerate(rows_data):
        c0, c1 = table.row_cells(i)
        set_cell_bg(c0, '1E2E40')
        set_cell_bg(c1, 'FAFAFA')
        r0 = c0.paragraphs[0].add_run(campo)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = GOLD
        r1 = c1.paragraphs[0].add_run(valor)
        r1.font.size = Pt(9.5)
        for cell in [c0, c1]:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    add_border(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '1E2E40')
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GOLD
        hdr[i].paragraphs[0].paragraph_format.space_before = Pt(2)
        hdr[i].paragraphs[0].paragraph_format.space_after  = Pt(2)
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        bg = 'F7F9FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            run = cells[ci].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            cells[ci].paragraphs[0].paragraph_format.space_before = Pt(2)
            cells[ci].paragraphs[0].paragraph_format.space_after  = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    doc.add_paragraph()

def bullet(doc, text, level=0):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after   = Pt(2)
    return p

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(3)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SISTEMA DE GESTIÓN')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complejo Riviera Maya')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = GOLD

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Salón de Eventos')
run.font.size = Pt(16)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

hr(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Técnica del Sistema')
run.font.size = Pt(13)
run.font.color.rgb = DGRAY

doc.add_paragraph()

info = [
    ('Versión',  '1.0'),
    ('Fecha',    'Junio 2025'),
    ('Sistema',  'Django 6 · Python 3.12 · SQLite / PostgreSQL'),
    ('Autores',  'Equipo de Desarrollo'),
]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK
    r2 = p.add_run(val)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, 'Índice de Contenido')

toc = [
    ('1.', 'Diagramas de Casos de Uso'),
    ('  1.1.', 'Diagrama General del Sistema'),
    ('  1.2.', 'Subsistema de Cotización (detalle)'),
    ('  1.3.', 'Subsistema de Operaciones — Staff'),
    ('2.', 'Especificaciones de Casos de Uso'),
    ('  UC-01', 'Completar Proceso de Cotización y Reserva'),
    ('  UC-02', 'Liquidar Saldo Pendiente de Reserva'),
    ('  UC-03', 'Registrar Pago Manual (Staff)'),
    ('  UC-04', 'Consultar Dashboard de Operaciones'),
    ('  UC-05', 'Gestionar Agenda Mensual (Staff)'),
    ('3.', 'Mock-ups del Sistema'),
    ('  3.1.', 'Página Principal (Home)'),
    ('  3.2.', 'Catálogo de Espacios'),
    ('  3.3.', 'Cotizador — Wizard 7 Pasos'),
    ('  3.4.', 'Portal Cliente — Mis Reservas'),
    ('  3.5.', 'Panel de Operaciones — Dashboard'),
    ('  3.6.', 'Panel de Operaciones — Agenda Mensual'),
    ('  3.7.', 'Panel de Operaciones — Detalle de Reservación'),
    ('4.', 'Modelado del Sistema'),
    ('  4.1.', 'Diagramas de Secuencia'),
    ('  4.2.', 'Diagramas de Actividades'),
    ('  4.3.', 'Diagrama de Clases'),
    ('  4.4.', 'Diagrama de Base de Datos (ERD)'),
]
for num, title in toc:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = GOLD
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  1. DIAGRAMAS DE CASOS DE USO
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '1. Diagramas de Casos de Uso')
note_box(doc,
    'Nota: Los diagramas de casos de uso están especificados en lenguaje PlantUML. '
    'Para renderizarlos visualmente pega el código en https://www.plantuml.com/plantuml '
    'o usa la extensión PlantUML en VS Code (Alt+D con el cursor dentro del bloque).')

# ── 1.1 General ─────────────────────────────────────────────────────────────
heading2(doc, '1.1 Diagrama General del Sistema')

body(doc, 'El sistema cuenta con 4 actores principales y 32 casos de uso distribuidos en 6 subsistemas:')

simple_table(doc,
    ['Actor', 'Descripción', 'Acceso'],
    [
        ['Visitante', 'Usuario no autenticado. Puede explorar el catálogo e iniciar cotización.', 'Público'],
        ['Cliente Registrado', 'Extiende Visitante. Accede al portal personal para gestionar sus reservas.', 'Login requerido'],
        ['Staff / Admin', 'Personal operativo. Accede al panel de administración y reportes.', 'is_staff=True'],
        ['Sistema', 'Automatizaciones internas que se ejecutan al confirmar una reserva.', 'Interno'],
    ],
    [Cm(3.5), Cm(9), Cm(3)]
)

body(doc, 'Resumen de casos de uso por paquete:')

simple_table(doc,
    ['Paquete', 'Casos de Uso', 'Actor(es)'],
    [
        ['Catálogo y Disponibilidad',      'UC01–UC04  Ver espacios, detalles, disponibilidad y paquetes', 'Visitante'],
        ['Cotización (Wizard 7 pasos)',    'UC05–UC11  Proceso completo: tipo → espacio → fecha → paquete → extras → datos → pago', 'Visitante'],
        ['Portal Cliente',                 'UC12–UC15  Ver reservas, detalle, liquidar saldo, ver contrato', 'Cliente Registrado'],
        ['Autenticación',                  'UC16–UC17  Iniciar y cerrar sesión', 'Visitante / Staff'],
        ['Operaciones (Staff)',             'UC18–UC26  Dashboard, agenda, reservas, pagos, reportes, tareas', 'Staff'],
        ['Automatización del Sistema',     'UC27–UC32  Generar código, crear usuario, contraseña, bloquear fecha, anticipo, bitácora', 'Sistema'],
    ],
    [Cm(4), Cm(8), Cm(3.5)]
)

code_block(doc, '''@startuml UCG_General
left to right direction
skinparam packageStyle rectangle
skinparam actorStyle awesome

actor "Visitante" as V
actor "Cliente\\nRegistrado" as C
actor "Staff / Admin" as S
actor "Sistema" as SYS

V <|-- C

rectangle "Complejo Riviera Maya — Sistema de Reservas" {

  package "Catálogo y Disponibilidad" {
    usecase "UC01 Ver catálogo\\nde espacios" as UC01
    usecase "UC02 Ver detalle\\nde espacio" as UC02
    usecase "UC03 Consultar\\ndisponibilidad" as UC03
    usecase "UC04 Ver paquetes\\ny precios" as UC04
  }

  package "Cotización (Wizard 7 pasos)" {
    usecase "UC05 Paso 1 Tipo, guests y hora" as UC05
    usecase "UC06 Paso 2 Seleccionar espacio" as UC06
    usecase "UC07 Paso 3 Seleccionar fecha" as UC07
    usecase "UC08 Paso 4 Seleccionar paquete" as UC08
    usecase "UC09 Paso 5 Servicios adicionales" as UC09
    usecase "UC10 Paso 6 Datos del cliente" as UC10
    usecase "UC11 Paso 7 Procesar pago (25%)" as UC11
  }

  package "Portal Cliente" {
    usecase "UC12 Ver mis reservas" as UC12
    usecase "UC13 Ver detalle de reserva" as UC13
    usecase "UC14 Liquidar saldo pendiente" as UC14
    usecase "UC15 Ver contrato" as UC15
  }

  package "Operaciones (Staff)" {
    usecase "UC18 Ver dashboard y KPIs" as UC18
    usecase "UC19 Gestionar agenda mensual" as UC19
    usecase "UC20 Listar reservaciones" as UC20
    usecase "UC22 Registrar pago manual" as UC22
    usecase "UC23 Cambiar estado de reserva" as UC23
    usecase "UC25 Ver reportes anuales" as UC25
    usecase "UC26 Gestionar tareas" as UC26
  }

  package "Automatización del Sistema" {
    usecase "UC27 Generar código de reserva" as UC27
    usecase "UC28 Crear usuario automáticamente" as UC28
    usecase "UC29 Generar contraseña temporal" as UC29
    usecase "UC30 Bloquear fecha ocupada" as UC30
    usecase "UC31 Registrar anticipo (25%)" as UC31
    usecase "UC32 Registrar actividad en bitácora" as UC32
  }
}

V --> UC01 \nV --> UC11 \nC --> UC12 \nC --> UC14
S --> UC18 \nS --> UC22 \nSYS --> UC27 \nSYS --> UC32

UC11 ..> UC27 : <<include>>
UC11 ..> UC28 : <<include>>
UC11 ..> UC30 : <<include>>
UC11 ..> UC31 : <<include>>
UC14 ..> UC22 : <<extend>>
UC22 ..> UC32 : <<include>>
@enduml''', title='Código PlantUML — Diagrama General (pegar en plantuml.com):')

# ── 1.2 Cotización ───────────────────────────────────────────────────────────
heading2(doc, '1.2 Subsistema de Cotización (Detalle)')
body(doc, 'Muestra el flujo interno del wizard de 7 pasos con sus relaciones include y extend:')

code_block(doc, '''@startuml UCS_Cotizacion
left to right direction
actor "Visitante" as V
actor "Sistema" as SYS

rectangle "Subsistema de Cotización" {
  usecase "Iniciar cotización (/cotizar/)" as UC_INIT
  usecase "Paso 1: Tipo, guests y hora" as UC_P1
  usecase "Paso 2: Seleccionar espacio" as UC_P2
  usecase "Paso 3: Seleccionar fecha" as UC_P3
  usecase "Verificar disponibilidad (FechaOcupada)" as UC_DISP
  usecase "Paso 4: Seleccionar paquete" as UC_P4
  usecase "Paso 5: Servicios adicionales" as UC_P5
  usecase "Paso 6: Datos personales" as UC_P6
  usecase "Paso 7: Procesar pago anticipo 25%" as UC_P7
  usecase "Calcular total y anticipo" as UC_CALC
  usecase "Crear reserva con todos sus registros" as UC_CREATE
  usecase "Mostrar confirmación y contraseña" as UC_CONFIRM
}

V --> UC_INIT
UC_INIT ..> UC_P1 : <<include>>
UC_P1 ..> UC_P2 : <<include>>
UC_P2 ..> UC_P3 : <<include>>
UC_P3 ..> UC_DISP : <<include>>
UC_DISP ..> UC_P4 : <<extend>> [si fecha disponible]
UC_P4 ..> UC_P7 : <<include>>
UC_P7 ..> UC_CALC : <<include>>
UC_P7 ..> UC_CREATE : <<include>>
UC_CREATE ..> UC_CONFIRM : <<include>>
SYS --> UC_DISP
SYS --> UC_CREATE
@enduml''', title='Código PlantUML — Subsistema de Cotización:')

# ── 1.3 Operaciones ──────────────────────────────────────────────────────────
heading2(doc, '1.3 Subsistema de Operaciones (Staff)')

simple_table(doc,
    ['Grupo', 'Casos de Uso'],
    [
        ['Dashboard',           'Ver KPIs del mes, próximos eventos (7d), ocupación por espacio, reservas por liquidar'],
        ['Agenda',              'Ver calendario mensual, filtrar por espacio, ver detalle de día'],
        ['Gestión de Reservas', 'Listar con filtros, ver detalle, registrar pago manual, cambiar estado, notas internas'],
        ['Catálogo Operativo',  'Ver todos los espacios, paquetes y servicios activos'],
        ['Reportes',            'Ingresos anuales (gráfica mensual), distribución por tipo de evento, top espacios'],
    ],
    [Cm(4.5), Cm(11)]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  2. ESPECIFICACIONES DE CASOS DE USO
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '2. Especificaciones de Casos de Uso')

# ── UC-01 ────────────────────────────────────────────────────────────────────
heading2(doc, 'UC-01 — Completar Proceso de Cotización y Reserva')

uc_table(doc, [
    ('ID',                 'UC-01'),
    ('Nombre',             'Completar Proceso de Cotización y Reserva'),
    ('Actor principal',    'Visitante (puede ser nuevo cliente o cliente existente)'),
    ('Actores secundarios','Sistema'),
    ('Precondiciones',     'Ninguna — acceso público sin requerir sesión activa'),
    ('Postcondiciones',    'Reserva creada en BD, fecha bloqueada, cliente autenticado, anticipo registrado'),
    ('Prioridad',          'Alta'),
    ('Frecuencia de uso',  'Varias veces por semana'),
])

heading3(doc, 'Flujo Principal de Eventos')

steps = [
    ('1',  'Visitante', 'Accede a /cotizar/'),
    ('2',  'Sistema',   'Muestra formulario Paso 1 con opciones de tipo de evento, campo de invitados y selector de hora'),
    ('3',  'Visitante', 'Selecciona tipo de evento (boda/XV/bautizo/aniversario/empresarial/cumpleaños), ingresa invitados (1-600) y hora de inicio (14:00-20:00)'),
    ('4',  'Sistema',   'Valida campos; guarda en request.session["cotizador"]; redirige al Paso 2'),
    ('5',  'Visitante', 'Selecciona un espacio de las tarjetas disponibles'),
    ('6',  'Sistema',   'Guarda espacio_id en sesión; actualiza resumen de costo en sidebar; redirige a Paso 3'),
    ('7',  'Visitante', 'Selecciona una fecha en el calendario del evento'),
    ('8',  'Sistema',   'Consulta FechaOcupada para ese espacio; si disponible, guarda fecha y redirige a Paso 4'),
    ('9',  'Visitante', 'Selecciona paquete: básico / premium / lujo'),
    ('10', 'Sistema',   'Guarda paquete_id; actualiza total estimado en sidebar; redirige a Paso 5'),
    ('11', 'Visitante', 'Selecciona servicios adicionales (múltiple, opcional)'),
    ('12', 'Sistema',   'Guarda servicios[] en sesión; recalcula total; redirige a Paso 6'),
    ('13', 'Visitante', 'Ingresa nombre, apellidos, email y teléfono'),
    ('14', 'Sistema',   'Guarda datos de contacto en sesión; redirige a Paso 7'),
    ('15', 'Visitante', 'Ingresa titular, número tarjeta (16 dígitos), vencimiento MM/AA, CVV (3 dígitos) y acepta términos'),
    ('16', 'Sistema',   'Valida formato de todos los campos del formulario de pago'),
    ('17', 'Sistema',   'Calcula: total = precio_espacio + (precio_paquete × guests) + extras; anticipo = total × 0.25'),
    ('18', 'Sistema',   'User.objects.get_or_create(username=email): si nuevo, genera contraseña 10 chars y crea ClienteProfile'),
    ('19', 'Sistema',   'Crea Reserva con código único RM-{YYYY}-{4 dígitos} y todos los campos del wizard'),
    ('20', 'Sistema',   'Crea registros ReservaServicio por cada servicio adicional seleccionado'),
    ('21', 'Sistema',   'Crea FechaOcupada(espacio, fecha, motivo="reserva") — bloquea la fecha para ese espacio'),
    ('22', 'Sistema',   'Crea PagoRegistro(monto=anticipo, concepto="anticipo")'),
    ('23', 'Sistema',   'Crea ActividadReserva(descripcion="Reserva creada vía cotizador web")'),
    ('24', 'Sistema',   'Ejecuta auth.login(request, user); limpia session["cotizador"]'),
    ('25', 'Sistema',   'Redirige a /cotizar/confirmacion/{codigo}/ con resumen y contraseña temporal si es nuevo usuario'),
]
for num, actor, action in steps:
    flow_step(doc, num, actor, action)
doc.add_paragraph()

heading3(doc, 'Flujos Alternativos')

alt_flows = [
    ('FA-1: Fecha no disponible (Paso 8)',
     ['Sistema detecta que FechaOcupada contiene esa combinación espacio+fecha',
      'Muestra error: "Lo sentimos, esta fecha ya no está disponible para el espacio seleccionado"',
      'Visitante elige otra fecha (regresa al Paso 7)']),
    ('FA-2: Usuario ya registrado (Paso 18)',
     ['get_or_create retorna created=False',
      'Sistema usa el User existente sin alterar su contraseña',
      'Continúa desde Paso 19']),
    ('FA-3: Datos de tarjeta inválidos (Paso 16)',
     ['Validación falla: tarjeta ≠ 16 dígitos, vencimiento expirado, CVV ≠ 3 dígitos, o sin aceptar términos',
      'Re-renderiza el formulario Paso 7 con los errores indicados por campo',
      'Visitante corrige y reenvía']),
    ('FA-4: Sesión expirada o incompleta',
     ['Sistema detecta que request.session["cotizador"] está vacío o falta algún paso previo',
      'Redirige al Paso 1 con mensaje informativo']),
]
for title, items in alt_flows:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    for item in items:
        bullet(doc, item)
    doc.add_paragraph()

heading3(doc, 'Reglas de Negocio')
rules = [
    'Anticipo obligatorio: exactamente el 25% del total del evento',
    'Invitados: mínimo 1, máximo 600 personas',
    'Horario de inicio: de 14:00 a 20:00 hrs (7 opciones, cada hora)',
    'Código de reserva: formato RM-{YYYY}-{4 dígitos}, único e irrepetible',
    'Una fecha bloquea el espacio completo (un evento por día por espacio según eventos_por_dia)',
    'La contraseña temporal se muestra una sola vez en la página de confirmación',
]
for r in rules:
    bullet(doc, r)

doc.add_paragraph()

# ── UC-02 ────────────────────────────────────────────────────────────────────
heading2(doc, 'UC-02 — Liquidar Saldo Pendiente de Reserva')

uc_table(doc, [
    ('ID',              'UC-02'),
    ('Nombre',          'Liquidar Saldo Pendiente de Reserva'),
    ('Actor principal', 'Cliente Registrado'),
    ('Precondiciones',  'Cliente autenticado; reserva con estado = confirmada o pendiente_liquidacion'),
    ('Postcondiciones', 'PagoRegistro creado con concepto liquidacion; anticipo_pagado actualizado; estado cambia a completada'),
    ('Prioridad',       'Alta'),
])

heading3(doc, 'Flujo Principal')
for num, actor, action in [
    ('1', 'Cliente', 'Accede a "Mis Reservas" en /portal/reservas/'),
    ('2', 'Sistema', 'Muestra listado de reservas del cliente con el saldo pendiente resaltado'),
    ('3', 'Cliente', 'Hace clic en "Liquidar saldo" en la reserva deseada'),
    ('4', 'Sistema', 'Calcula saldo = total_evento − anticipo_pagado; muestra LiquidacionForm con el monto'),
    ('5', 'Cliente', 'Ingresa datos de tarjeta y marca "Confirmar pago"'),
    ('6', 'Sistema', 'Valida LiquidacionForm (mismas reglas que Paso 7 del wizard)'),
    ('7', 'Sistema', 'Crea PagoRegistro(monto=saldo, concepto="liquidacion")'),
    ('8', 'Sistema', 'Actualiza reserva.anticipo_pagado += saldo'),
    ('9', 'Sistema', 'Cambia reserva.estado = "completada"'),
    ('10','Sistema', 'Crea ActividadReserva("Liquidación completada por el cliente")'),
    ('11','Sistema', 'Redirige a /portal/reservas/{codigo}/ con banner de confirmación'),
]:
    flow_step(doc, num, actor, action)
doc.add_paragraph()

# ── UC-03 ────────────────────────────────────────────────────────────────────
heading2(doc, 'UC-03 — Registrar Pago Manual (Staff)')

uc_table(doc, [
    ('ID',              'UC-03'),
    ('Nombre',          'Registrar Pago Manual en Reservación'),
    ('Actor principal', 'Staff / Admin'),
    ('Precondiciones',  'Usuario staff autenticado (is_staff=True); reserva existente en el sistema'),
    ('Postcondiciones', 'PagoRegistro creado; anticipo_pagado actualizado; actividad registrada en bitácora'),
    ('Prioridad',       'Media'),
])

heading3(doc, 'Flujo Principal')
for num, actor, action in [
    ('1', 'Staff',   'Accede a /ops/reservas/{codigo}/'),
    ('2', 'Sistema', 'Verifica @staff_member_required; carga reserva con servicios, pagos y bitácora'),
    ('3', 'Staff',   'Localiza el panel "Registrar pago" en la vista de detalle'),
    ('4', 'Sistema', 'Muestra formulario: monto, fecha_pago, concepto (anticipo / liquidacion / extra)'),
    ('5', 'Staff',   'Completa el formulario y envía'),
    ('6', 'Sistema', 'Crea PagoRegistro(monto, fecha_pago, concepto, registrado_por=staff_user)'),
    ('7', 'Sistema', 'Actualiza reserva.anticipo_pagado += monto; reserva.save()'),
    ('8', 'Sistema', 'Crea ActividadReserva(f"Pago ${monto} registrado por {staff}", usuario=staff)'),
    ('9', 'Sistema', 'Redirige al detalle con el nuevo pago visible en el historial'),
]:
    flow_step(doc, num, actor, action)
doc.add_paragraph()

# ── UC-04 ────────────────────────────────────────────────────────────────────
heading2(doc, 'UC-04 — Consultar Dashboard de Operaciones')

uc_table(doc, [
    ('ID',              'UC-04'),
    ('Nombre',          'Consultar Dashboard y KPIs del Sistema'),
    ('Actor principal', 'Staff / Admin'),
    ('Precondiciones',  'Usuario con is_staff=True autenticado'),
    ('Postcondiciones', 'Ninguna — operación de solo lectura'),
    ('Prioridad',       'Media'),
])

heading3(doc, 'Flujo Principal')
for num, actor, action in [
    ('1', 'Staff',   'Accede a /ops/'),
    ('2', 'Sistema', 'Calcula eventos_mes: reservas del mes actual con estado confirmada o completada'),
    ('3', 'Sistema', 'Suma ingresos_confirmados de reservas del mes'),
    ('4', 'Sistema', 'Suma pipeline_total de todas las reservas activas'),
    ('5', 'Sistema', 'Obtiene proximos_eventos: reservas con fecha en los próximos 7 días'),
    ('6', 'Sistema', 'Calcula ocupacion_por_espacio: % de días con evento en el mes actual'),
    ('7', 'Sistema', 'Obtiene reservas_liquidar: reservas con saldo pendiente y evento en ≤ 7 días'),
    ('8', 'Sistema', 'Lista tareas_pendientes sin completar (TareaPendiente.filter(completada=False))'),
    ('9', 'Sistema', 'Renderiza dashboard completo con todos los datos calculados'),
]:
    flow_step(doc, num, actor, action)
doc.add_paragraph()

# ── UC-05 ────────────────────────────────────────────────────────────────────
heading2(doc, 'UC-05 — Gestionar Agenda Mensual (Staff)')

uc_table(doc, [
    ('ID',              'UC-05'),
    ('Nombre',          'Gestionar Agenda y Calendario Mensual'),
    ('Actor principal', 'Staff / Admin'),
    ('Precondiciones',  'Usuario staff autenticado'),
    ('Postcondiciones', 'Ninguna — visualización de solo lectura'),
    ('Prioridad',       'Media'),
])

heading3(doc, 'Flujo Principal')
for num, actor, action in [
    ('1', 'Staff',   'Accede a /ops/agenda/'),
    ('2', 'Sistema', 'Obtiene mes/año actual (o del parámetro GET); construye estructura de calendario'),
    ('3', 'Sistema', 'Carga todas las reservas del mes y las asocia a sus días correspondientes'),
    ('4', 'Sistema', 'Calcula métricas: total eventos, invitados, ingreso esperado, días ocupados'),
    ('5', 'Sistema', 'Renderiza calendario con días resaltados donde hay eventos activos'),
    ('6', 'Staff',   'Aplica filtro por espacio (GET ?espacio=id)'),
    ('7', 'Sistema', 'Re-renderiza mostrando solo eventos del espacio seleccionado'),
    ('8', 'Staff',   'Hace clic en un día con evento'),
    ('9', 'Sistema', 'Muestra panel con detalle del evento: código, cliente, tipo, invitados, monto'),
]:
    flow_step(doc, num, actor, action)
doc.add_paragraph()

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  3. MOCK-UPS
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '3. Mock-ups del Sistema')
note_box(doc, 'Los mock-ups utilizan notación ASCII para representar la estructura visual de cada pantalla. '
              'Representan el layout, jerarquía de información y elementos de interacción sin especificar estilos finales.')

# ── 3.1 Home ─────────────────────────────────────────────────────────────────
heading2(doc, '3.1 Página Principal (Home)')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  ⚡ RIVIERA MAYA                        [Iniciar sesión]  [Cotizar] │
│  [Inicio] [Espacios] [Paquetes] [Disponibilidad]                    │
├──────────────────────────────────────────────────────────────────────┤
│                                                                      │
│  ╔════════════════════════════════════════════════════╗              │
│  ║       TU EVENTO PERFECTO TE ESPERA                ║              │
│  ║   Espacios únicos para momentos inolvidables       ║              │
│  ║            [ COTIZAR MI EVENTO ]                   ║              │
│  ╚════════════════════════════════════════════════════╝              │
│                    — IMAGEN HERO —                                   │
│                                                                      │
├──────────────────────────────────────────────────────────────────────┤
│  NUESTROS ESPACIOS DESTACADOS                                        │
│  ┌────────────┐  ┌────────────┐  ┌────────────┐  ┌────────────┐   │
│  │ [IMAGEN]   │  │ [IMAGEN]   │  │ [IMAGEN]   │  │ [IMAGEN]   │   │
│  │ Jardín     │  │ Salón      │  │ Terraza    │  │ Palapa     │   │
│  │ Tropical   │  │ Imperial   │  │ del Mar    │  │ Colonial   │   │
│  │ 300 pers.  │  │ 500 pers.  │  │ 200 pers.  │  │ 150 pers.  │   │
│  │ $15,000+   │  │ $22,000+   │  │ $18,000+   │  │ $12,000+   │   │
│  │[Ver espacio│  │[Ver espacio│  │[Ver espacio│  │[Ver espacio│   │
│  └────────────┘  └────────────┘  └────────────┘  └────────────┘   │
├──────────────────────────────────────────────────────────────────────┤
│  NUESTROS PAQUETES                                                   │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐             │
│  │   BÁSICO     │  │ ★ PREMIUM ★  │  │    LUJO      │             │
│  │ ✓ Mesas      │  │ ✓ Todo bás.  │  │ ✓ Todo prem. │             │
│  │ ✓ Mantelería │  │ ✓ Decoración │  │ ✓ Chef priv. │             │
│  │ ✓ Iluminación│  │ ✓ DJ/música  │  │ ✓ Open bar   │             │
│  │  [ Cotizar ] │  │  [ Cotizar ] │  │  [ Cotizar ] │             │
│  └──────────────┘  └──────────────┘  └──────────────┘             │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.2 Catálogo ─────────────────────────────────────────────────────────────
heading2(doc, '3.2 Catálogo de Espacios')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  CATÁLOGO DE ESPACIOS                                                │
│  Filtrar: [Todos] [Interior] [Exterior] [Vista al mar]              │
│  Capacidad: [Hasta 200] [200-400] [Más de 400]                      │
│                                                                      │
│  ┌────────────────────────┐  ┌────────────────────────┐            │
│  │   [IMAGEN 400×280]     │  │   [IMAGEN 400×280]     │            │
│  │ Jardín Tropical        │  │ Salón Imperial         │            │
│  │ 🌿 Exterior Tropical   │  │ ❄️ Interior Climatiz.  │            │
│  │ 👥 Hasta 300 personas  │  │ 👥 Hasta 500 personas  │            │
│  │ 💰 Desde $15,000       │  │ 💰 Desde $22,000       │            │
│  │         [Ver detalles] │  │         [Ver detalles] │            │
│  └────────────────────────┘  └────────────────────────┘            │
│                                                                      │
│  ┌────────────────────────┐  ┌────────────────────────┐            │
│  │   [IMAGEN 400×280]     │  │   [IMAGEN 400×280]     │            │
│  │ Terraza del Mar        │  │ Palapa Colonial        │            │
│  │ 🌊 Exterior+Vista Mar  │  │ 🏛️ Interior Versátil  │            │
│  │ 👥 Hasta 200 personas  │  │ 👥 Hasta 150 personas  │            │
│  │ 💰 Desde $18,000       │  │ 💰 Desde $12,000       │            │
│  │         [Ver detalles] │  │         [Ver detalles] │            │
│  └────────────────────────┘  └────────────────────────┘            │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.3 Wizard ───────────────────────────────────────────────────────────────
heading2(doc, '3.3 Cotizador — Wizard 7 Pasos')
heading3(doc, 'Paso 1 — Tipo de Evento')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  COTIZAR MI EVENTO — Paso 1 de 7: Detalles del evento               │
│  ● ──── ○ ──── ○ ──── ○ ──── ○ ──── ○ ──── ○                       │
│  [1]   [2]   [3]   [4]   [5]   [6]   [7]                           │
│  Tipo  Esp.  Fech. Paq.  Ext.  Dat.  Pago                          │
├────────────────────────────────┬─────────────────────────────────────┤
│                                │  📋 RESUMEN COTIZACIÓN              │
│  ¿Cuál es tu evento?           │  Tipo:          —                   │
│  ○ 💍 Boda                    │  Invitados:     —                   │
│  ○ 🌹 XV Años                 │  Hora:          —                   │
│  ○ 🕊️ Bautizo               │  Espacio:       —                   │
│  ○ 💑 Aniversario             │  Paquete:       —                   │
│  ○ 💼 Empresarial             │  ─────────────────────────────      │
│  ○ 🎂 Cumpleaños              │  Subtotal:  $ ——,———                │
│                                │  Anticipo:  $ ——,———                │
│  Invitados: [  150  ]          │                                     │
│  (mín 1, máx 600)              │                                     │
│  Hora: [ 17:00 hrs ▼ ]         │                                     │
│            [ Continuar → ]     │                                     │
└────────────────────────────────┴─────────────────────────────────────┘''')

heading3(doc, 'Paso 3 — Selección de Fecha')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  COTIZAR MI EVENTO — Paso 3 de 7: Selección de fecha               │
│  ✓ ──── ✓ ──── ● ──── ○ ──── ○ ──── ○ ──── ○                       │
├────────────────────────────────┬─────────────────────────────────────┤
│  Elige la fecha de tu evento   │  Tipo:       Boda                   │
│                                │  Invitados:  150                    │
│        JUNIO 2025              │  Hora:       17:00 hrs              │
│  ┌──────────────────────────┐  │  Espacio:    Jardín Tropical        │
│  │ Lu  Ma  Mi  Ju  Vi  Sá Do│  │  ─────────────────────────────      │
│  │                  1   2  │  │  Subtotal:    $18,000               │
│  │  3   4   5   6   7  🔴  9 │  │  Anticipo:    $ ——,———             │
│  │ 10  11  12  13  14  15 16 │  │                                     │
│  │ 17  18  19  20  21 [22]23 │  │  🔴 Fecha ocupada                  │
│  │ 24  25  26  27  28  🔴 30 │  │  🟢 Disponible                     │
│  └──────────────────────────┘  │  🔵 Seleccionada                    │
│  Fecha: [ 22/06/2025 ]         │                                     │
│  [← Volver]  [ Continuar → ]  │                                     │
└────────────────────────────────┴─────────────────────────────────────┘''')

heading3(doc, 'Paso 7 — Pago (Anticipo 25%)')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  COTIZAR MI EVENTO — Paso 7 de 7: Pago                             │
│  ✓ ──── ✓ ──── ✓ ──── ✓ ──── ✓ ──── ✓ ──── ●                       │
├────────────────────────────────┬─────────────────────────────────────┤
│  Datos de pago                 │  RESUMEN FINAL                      │
│                                │  Jardín Tropical:  $15,000          │
│  Titular: [Juan García López ] │  Paquete Premium:                   │
│  Tarjeta: [···· ···· ···· ····]│    150 × $350:     $52,500          │
│  Venct.:  [MM/AA]  CVV: [···]  │  DJ Profesional:   $ 5,000          │
│                                │  Foto y video:     $ 3,500          │
│  ☐ Acepto términos y           │  ─────────────────────────────      │
│    condiciones                 │  TOTAL EVENTO:     $76,000          │
│                                │                                     │
│  [← Volver]                    │  Anticipo (25%):   $19,000          │
│  [ ✓ CONFIRMAR Y PAGAR ]       │  Saldo restante:   $57,000          │
└────────────────────────────────┴─────────────────────────────────────┘''')

heading3(doc, 'Confirmación de Reserva')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│                    ✅ ¡RESERVA CONFIRMADA!                           │
│           Tu código de reserva es:  RM-2025-0047                    │
├──────────────────────────────────────────────────────────────────────┤
│  Evento:      Boda · 150 invitados                                   │
│  Espacio:     Jardín Tropical                                        │
│  Fecha:       Sábado 22 de Junio 2025 · 17:00 hrs                  │
│  Paquete:     Premium                                                │
│  Total:       $76,000 │ Anticipo pagado: $19,000 │ Saldo: $57,000   │
├──────────────────────────────────────────────────────────────────────┤
│  🔑 Tu cuenta fue creada automáticamente:                           │
│     Usuario:     juan@email.com                                      │
│     Contraseña:  Xk9mP2qA7n  (cámbiala al iniciar sesión)          │
├──────────────────────────────────────────────────────────────────────┤
│       [ Ver mis reservas ]       [ Descargar contrato ]             │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.4 Portal cliente ────────────────────────────────────────────────────────
heading2(doc, '3.4 Portal Cliente — Mis Reservas')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  RIVIERA MAYA · Mi Portal          [Hola, Juan García ▼]  [Salir]  │
├──────────────────────────────────────────────────────────────────────┤
│  MIS RESERVAS                                                        │
│  ┌────────────┐  ┌──────────────────────┐  ┌──────────────────┐   │
│  │ Activas: 2 │  │ Total invertido      │  │ Próximo evento   │   │
│  │            │  │   $95,000            │  │ 22 Jun 2025      │   │
│  └────────────┘  └──────────────────────┘  └──────────────────┘   │
│  [Todas] [Confirmadas] [Pendiente pago] [Completadas]               │
│                                                                      │
│  ┌──────────────────────────────────────────────────────────────┐   │
│  │  RM-2025-0047                           ● CONFIRMADA         │   │
│  │  Jardín Tropical · Boda · Sáb 22 Jun 2025 · 17:00 hrs       │   │
│  │  150 invitados · Paquete Premium                             │   │
│  │  Total: $76,000 │ Anticipo: $19,000  ████░░░░░  25%          │   │
│  │  Saldo pendiente: $57,000                                    │   │
│  │                   [ Ver contrato ]  [ Liquidar saldo ]       │   │
│  └──────────────────────────────────────────────────────────────┘   │
│                                                                      │
│  ┌──────────────────────────────────────────────────────────────┐   │
│  │  RM-2025-0021                         ✓ COMPLETADA           │   │
│  │  Salón Imperial · Aniversario · Sáb 10 Mar 2025 · 18:00     │   │
│  │  80 invitados · Paquete Lujo                                 │   │
│  │  Total: $48,500 │ Pagado: $48,500  ████████████████  100%   │   │
│  │                                         [ Ver contrato ]     │   │
│  └──────────────────────────────────────────────────────────────┘   │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.5 Dashboard ─────────────────────────────────────────────────────────────
heading2(doc, '3.5 Panel de Operaciones — Dashboard')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  ≡ PANEL DE OPERACIONES                      [Admin ▼]  [Salir]    │
│  [Dashboard] [Agenda] [Reservas] [Catálogo] [Reportes]              │
├──────────────────────────────────────────────────────────────────────┤
│  Buenos días, Administrador · Junio 2025                            │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐  ┌────────┐ │
│  │ Eventos/Mes  │  │ Ingresos     │  │ Pipeline     │  │  Ocup. │ │
│  │     12       │  │  $145,000    │  │  $280,500    │  │  72%   │ │
│  └──────────────┘  └──────────────┘  └──────────────┘  └────────┘ │
│  ┌────────────────────────────────┐  ┌────────────────────────┐   │
│  │  PRÓXIMOS EVENTOS (7 días)     │  │  TAREAS PENDIENTES     │   │
│  │  12 Jun · Salón Imperial       │  │  🔴 Confirmar RM-0045  │   │
│  │    Boda · 250 inv. RM-0042     │  │  🟡 Enviar cotización  │   │
│  │  13 Jun · Jardín Tropical      │  │  🟢 Actualizar fotos   │   │
│  │    XV Años · 180 inv. RM-0039  │  │  [ + Nueva tarea ]     │   │
│  └────────────────────────────────┘  └────────────────────────┘   │
│  ┌────────────────────────────────┐  ┌────────────────────────┐   │
│  │  OCUPACIÓN POR ESPACIO         │  │  POR LIQUIDAR (7 días) │   │
│  │  Salón Imperial  ██████████ 85%│  │  RM-0047: $57,000      │   │
│  │  Jardín Trop.    ████████  72% │  │  RM-0039: $38,000      │   │
│  │  Terraza Mar     ██████    60% │  │  RM-0041: $22,500      │   │
│  │  Palapa Col.     █████     45% │  │  Total:  $117,500      │   │
│  └────────────────────────────────┘  └────────────────────────┘   │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.6 Agenda ────────────────────────────────────────────────────────────────
heading2(doc, '3.6 Panel de Operaciones — Agenda Mensual')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  AGENDA · JUNIO 2025                  Filtrar: [Todos los espacios] │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐          │
│  │Eventos:12│  │Inv: 2150 │  │$145,000  │  │Días ocup:9│          │
│  └──────────┘  └──────────┘  └──────────┘  └──────────┘          │
│  ← Mayo                                             Julio →        │
│  ┌─────┬─────┬─────┬─────┬─────┬─────┬─────┐                     │
│  │ Lun │ Mar │ Mié │ Jue │ Vie │ Sáb │ Dom │                     │
│  ├─────┼─────┼─────┼─────┼─────┼─────┼─────┤                     │
│  │  3  │  4  │  5  │  6  │  7  │█ 8 █│  9  │  █ = día con evento │
│  │ 10  │ 11  │█12 █│█13 █│ 14  │ 15  │ 16  │                     │
│  │ 17  │ 18  │ 19  │ 20  │ 21  │█22 █│ 23  │                     │
│  │ 24  │ 25  │ 26  │ 27  │ 28  │█29 █│ 30  │                     │
│  └─────┴─────┴─────┴─────┴─────┴─────┴─────┘                     │
│  Clic en Sábado 22:                                                │
│  ┌──────────────────────────────────────────────────────────────┐  │
│  │  📅 Sábado 22 Jun 2025 · RM-2025-0047 · Juan García         │  │
│  │  Boda · 150 invitados · 17:00 · Jardín Tropical · $76,000   │  │
│  │                          [Ver detalle de reserva →]          │  │
│  └──────────────────────────────────────────────────────────────┘  │
└──────────────────────────────────────────────────────────────────────┘''')

# ── 3.7 Detalle reservación ───────────────────────────────────────────────────
heading2(doc, '3.7 Operaciones — Detalle de Reservación')
code_block(doc, '''┌──────────────────────────────────────────────────────────────────────┐
│  ← Volver       RM-2025-0047                         ● CONFIRMADA  │
│  ┌───────────────────────────┐  ┌────────────────────────────────┐ │
│  │ DATOS DEL EVENTO          │  │ RESUMEN FINANCIERO             │ │
│  │ Tipo:    Boda             │  │ Total:    $76,000              │ │
│  │ Cliente: Juan García      │  │ Anticipo: $19,000 ✓           │ │
│  │ Email:   juan@email.com   │  │ Saldo:    $57,000              │ │
│  │ Tel:     555-1234-567     │  │ ████░░░░░░░░░░░░░░░  25%      │ │
│  │ Espacio: Jardín Tropical  │  │  [ Registrar pago ]            │ │
│  │ Fecha:   22 Jun 2025      │  │  [ Marcar confirmada ]         │ │
│  │ Hora:    17:00 hrs        │  └────────────────────────────────┘ │
│  │ Guests:  150 personas     │                                      │
│  │ Paquete: Premium          │                                      │
│  └───────────────────────────┘                                      │
│  SERVICIOS: Foto y video $3,500 · DJ Profesional $5,000            │
│                                                                      │
│  HISTORIAL DE PAGOS                                                  │
│  Fecha       Concepto   Monto      Registrado por                    │
│  22/01/2025  Anticipo   $19,000    Sistema (cotizador)               │
│                                                                      │
│  NOTAS INTERNAS                                                      │
│  [ Cliente solicita mesa redonda y flores blancas...              ]  │
│                                          [ Guardar notas ]          │
│                                                                      │
│  BITÁCORA                                                            │
│  22 Jan 10:32 — Reserva creada vía cotizador web                    │
│  22 Jan 10:32 — Anticipo de $19,000 registrado automáticamente      │
└──────────────────────────────────────────────────────────────────────┘''')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  4. MODELADO DEL SISTEMA
# ════════════════════════════════════════════════════════════════════════════
heading1(doc, '4. Modelado del Sistema')

# ── 4.1 Secuencia ─────────────────────────────────────────────────────────────
heading2(doc, '4.1 Diagramas de Secuencia')
note_box(doc, 'Los diagramas de secuencia están en sintaxis Mermaid. Para renderizarlos: '
              'pégalos en https://mermaid.live o usa la extensión "Mermaid Preview" en VS Code.')

heading3(doc, '4.1.1 Flujo Completo de Cotización y Reserva')
body(doc, 'Este diagrama muestra la interacción completa entre el cliente, el navegador, las vistas Django, la sesión y la base de datos durante el proceso de cotización de 7 pasos.')
code_block(doc, '''sequenceDiagram
    participant C as Cliente/Visitante
    participant B as Navegador
    participant V as Vista Django
    participant S as Sesión
    participant DB as Base de Datos

    C->>B: GET /cotizar/
    B->>V: Request GET paso_1
    V->>B: Renderiza Paso 1 (tipo, guests, hora)

    C->>B: POST tipo_evento + num_invitados + hora_inicio
    B->>V: Request POST paso_1
    V->>V: Valida Paso1Form
    V->>S: session[cotizador] = {tipo, guests, hora}
    V->>B: Redirect /cotizar/espacio/

    Note over V,DB: Pasos 2-6: misma lógica de POST → validar → guardar en sesión → redirect

    C->>B: POST datos de tarjeta + aceptar_terminos
    B->>V: Request POST paso_7
    V->>V: Valida Paso7Form (tarjeta 16d, CVV 3d, MM/AA)
    V->>V: _calcular_total(): espacio + paquete x guests + extras
    V->>DB: User.objects.get_or_create(username=email)
    alt Usuario nuevo
        DB-->>V: created=True
        V->>V: _generar_contrasena() → pwd_tmp (10 chars)
        V->>DB: user.set_password(pwd_tmp); user.save()
        V->>DB: ClienteProfile.objects.create(user, telefono)
    else Usuario existente
        DB-->>V: created=False
    end
    V->>DB: Reserva.objects.create(codigo=RM-YYYY-XXXX, ...)
    loop Por cada servicio seleccionado
        V->>DB: ReservaServicio.objects.create(reserva, servicio)
    end
    V->>DB: FechaOcupada.objects.create(espacio, fecha, motivo=reserva)
    V->>DB: PagoRegistro.objects.create(monto=anticipo, concepto=anticipo)
    V->>DB: ActividadReserva.objects.create("Reserva creada vía cotizador")
    V->>V: auth.login(request, user)
    V->>S: del session[cotizador]
    V->>B: Redirect /cotizar/confirmacion/{codigo}/
    B-->>C: Código RM-YYYY-XXXX + contraseña temporal (si es nuevo)''', title='Mermaid — Flujo de Cotización:')

heading3(doc, '4.1.2 Liquidación de Saldo por Cliente')
code_block(doc, '''sequenceDiagram
    participant C as Cliente Registrado
    participant B as Navegador
    participant V as Vista Django
    participant DB as Base de Datos

    C->>B: GET /portal/reservas/{codigo}/liquidar/
    B->>V: Request GET liquidar_reserva
    V->>V: @login_required — verifica sesión activa
    V->>DB: Reserva.objects.get(codigo=codigo, cliente=request.user)
    DB-->>V: Reserva
    V->>V: saldo = total_evento - anticipo_pagado
    V->>B: Renderiza LiquidacionForm con saldo exacto
    B-->>C: Formulario con monto a pagar

    C->>B: POST datos tarjeta + confirmar_pago=True
    B->>V: Request POST liquidar_reserva
    V->>V: Valida LiquidacionForm
    alt Datos inválidos
        V->>B: Re-renderiza con errores
    else Datos válidos
        V->>DB: PagoRegistro.objects.create(monto=saldo, concepto=liquidacion)
        V->>DB: reserva.anticipo_pagado += saldo; reserva.save()
        V->>DB: reserva.estado = completada; reserva.save()
        V->>DB: ActividadReserva("Liquidación completada por el cliente")
        V->>B: Redirect /portal/reservas/{codigo}/
        B-->>C: Detalle de reserva con estado COMPLETADA
    end''', title='Mermaid — Liquidación de Saldo:')

heading3(doc, '4.1.3 Registro de Pago por Staff')
code_block(doc, '''sequenceDiagram
    participant S as Staff
    participant B as Navegador
    participant V as Vista Django
    participant DB as Base de Datos

    S->>B: GET /ops/reservas/{codigo}/
    B->>V: Request GET reservas_detalle
    V->>V: @staff_member_required — verifica is_staff=True
    V->>DB: Reserva.objects.select_related().get(codigo)
    DB-->>V: Reserva + servicios + pagos + actividad
    V->>B: Renderiza detalle con panel "Registrar pago"
    B-->>S: Vista completa con historial y formulario

    S->>B: POST monto + fecha_pago + concepto
    B->>V: Request POST (action=registrar_pago)
    V->>DB: PagoRegistro.objects.create(monto, fecha, concepto, registrado_por=staff)
    V->>DB: reserva.anticipo_pagado += monto; reserva.save()
    V->>DB: ActividadReserva(f"Pago ${monto} registrado por {staff}", usuario=staff)
    V->>B: Redirect /ops/reservas/{codigo}/
    B-->>S: Vista actualizada con nuevo pago en historial''', title='Mermaid — Registro de Pago Staff:')

heading3(doc, '4.1.4 Carga del Dashboard de Operaciones')
code_block(doc, '''sequenceDiagram
    participant S as Staff
    participant V as Vista Django
    participant DB as Base de Datos

    S->>V: GET /ops/
    V->>V: @staff_member_required

    par Consultas paralelas de KPIs
        V->>DB: Reserva.filter(fecha__month=hoy.month).count()
        DB-->>V: eventos_mes = 12
    and
        V->>DB: Reserva.filter(estado=confirmada).aggregate(Sum(total))
        DB-->>V: ingresos = $145,000
    and
        V->>DB: Reserva.filter(fecha__range=[hoy, hoy+7d])
        DB-->>V: proximos_eventos = 5
    and
        V->>DB: TareaPendiente.filter(completada=False)
        DB-->>V: tareas = 3
    end

    V->>V: Construye context con todos los datos
    V->>S: Renderiza dashboard completo''', title='Mermaid — Carga del Dashboard:')

# ── 4.2 Actividades ───────────────────────────────────────────────────────────
heading2(doc, '4.2 Diagramas de Actividades')
note_box(doc, 'Los diagramas de actividades están en sintaxis PlantUML. '
              'Para renderizarlos: pégalos en https://www.plantuml.com/plantuml')

heading3(doc, '4.2.1 Proceso Completo de Cotización')
code_block(doc, '''@startuml Act_Cotizacion
start
:Visitante accede a /cotizar/;

repeat
  :Paso 1: Tipo de evento, invitados (1-600), hora (14-20h);
repeat while (¿Campos válidos?) is (No)
-> Sí;

:Paso 2: Selecciona espacio del catálogo;
:Paso 3: Selecciona fecha en calendario;
:Sistema consulta FechaOcupada(espacio, fecha);

if (¿Fecha disponible?) then (No)
  #Pink:Muestra error "Fecha no disponible";
  goto Paso 3
else (Sí)
endif

:Paso 4: Selecciona paquete (básico/premium/lujo);
:Paso 5: Agrega servicios adicionales (opcional);
:Paso 6: Ingresa datos personales;

repeat
  :Paso 7: Datos de pago (tarjeta, MM/AA, CVV, términos);
repeat while (¿Datos válidos?) is (No)
-> Sí;

:Calcula total = espacio + paquete×guests + extras;
:Anticipo = total × 25%;
:Verifica email en base de datos;

if (¿Email registrado?) then (No)
  :Crea User, genera contraseña 10 chars, crea ClienteProfile;
else (Sí)
  :Recupera User existente;
endif

fork
  :Genera código RM-{YYYY}-{4d} · Crea Reserva;
fork again
  :Crea ReservaServicio por cada extra;
fork again
  :Crea FechaOcupada (bloquea fecha);
fork again
  :Crea PagoRegistro (anticipo 25%);
fork again
  :Crea ActividadReserva (log inicial);
end fork

:auth.login(request, user) · Limpia sesión;
:Muestra confirmación con código y contraseña temporal;
stop
@enduml''', title='PlantUML — Proceso de Cotización:')

heading3(doc, '4.2.2 Proceso de Autenticación y Redirección')
code_block(doc, '''@startuml Act_Autenticacion
start
:Usuario accede a /accounts/login/;
:Ingresa email y contraseña;
:Sistema ejecuta authenticate(username=email, password=pwd);

if (¿Credenciales válidas?) then (No)
  #Pink:Muestra error "Credenciales incorrectas";
  goto Sistema muestra LoginForm
else (Sí)
endif

:auth.login(request, user);

if (¿user.is_staff == True?) then (Sí)
  #LightBlue:Redirige a /ops/ (Dashboard de operaciones);
else (No)
  #LightGreen:Redirige a /portal/reservas/ (Portal del cliente);
endif
stop
@enduml''', title='PlantUML — Proceso de Autenticación:')

heading3(doc, '4.2.3 Proceso de Liquidación de Saldo')
code_block(doc, '''@startuml Act_Liquidacion
start
:Cliente accede a /portal/reservas/{codigo}/liquidar/;
:Sistema verifica @login_required;

if (¿Sesión activa?) then (No)
  :Redirige a /accounts/login/?next=...;
  stop
else (Sí)
endif

:Sistema obtiene Reserva (codigo + cliente=request.user);

if (¿Reserva existe y es del cliente?) then (No)
  #Pink:Error 404;
  stop
else (Sí)
endif

:Calcula saldo = total_evento - anticipo_pagado;
:Muestra LiquidacionForm con saldo exacto;

repeat
  :Cliente ingresa datos de tarjeta y confirma;
repeat while (¿LiquidacionForm válido?) is (No)
-> Sí;

:Crea PagoRegistro (liquidacion, monto=saldo);
:Actualiza reserva.anticipo_pagado += saldo;
:Cambia reserva.estado = completada;
:Crea ActividadReserva("Liquidación completada");
:Redirige al detalle con confirmación;
stop
@enduml''', title='PlantUML — Proceso de Liquidación:')

# ── 4.3 Clases ────────────────────────────────────────────────────────────────
heading2(doc, '4.3 Diagrama de Clases')
note_box(doc, 'Diagrama de clases completo del sistema con los 13 modelos Django, sus atributos, métodos y relaciones.')

# Tabla de clases
heading3(doc, 'Resumen de Clases del Sistema')
simple_table(doc,
    ['Clase', 'App', 'Atributos clave', 'Métodos / Propiedades'],
    [
        ['User', 'django.auth', 'id, username, email, first_name, last_name, is_staff, is_active', '(built-in Django)'],
        ['ClienteProfile', 'accounts', 'user (1:1), telefono, foto_perfil', '—'],
        ['Espacio', 'venues', 'nombre, slug, tipo, capacidad_max, precio_base, precio_max, estado, eventos_por_dia, orden', 'get_tipo_label(), get_tipo_grupo(), tiene_vista_al_mar()'],
        ['EspacioImagen', 'venues', 'espacio(FK), imagen, caption, orden', '—'],
        ['EspacioCaracteristica', 'venues', 'espacio(FK), descripcion, orden', '—'],
        ['Paquete', 'venues', 'nombre(choices), descripcion_corta, duracion_horas, precio_por_persona, es_mas_popular, orden', '—'],
        ['PaqueteItem', 'venues', 'paquete(FK), descripcion, incluido, orden', '—'],
        ['ServicioAdicional', 'venues', 'nombre, precio, precio_tipo, icono_nombre, activo, orden', 'precio_display()'],
        ['FechaOcupada', 'venues', 'espacio(FK), fecha, motivo(choices), nota | UNIQUE(espacio+fecha)', '—'],
        ['Reserva', 'bookings', 'codigo(unique), cliente(FK), espacio(FK), paquete(FK), tipo_evento, fecha, hora_inicio, num_invitados, estado, total_evento, anticipo_pagado, notas_internas, coordinador(FK)', 'saldo_pendiente(property), calcular_total(), get_estado_css()'],
        ['ReservaServicio', 'bookings', 'reserva(FK), servicio(FK), cantidad | UNIQUE(reserva+servicio)', 'subtotal()'],
        ['PagoRegistro', 'bookings', 'reserva(FK), monto, fecha_pago, concepto(choices), registrado_por(FK)', '—'],
        ['ActividadReserva', 'bookings', 'reserva(FK), descripcion, usuario(FK), created_at', '—'],
        ['TareaPendiente', 'operations', 'descripcion, completada, asignado_a(FK), prioridad(choices)', 'get_color()'],
    ],
    [Cm(3.5), Cm(2), Cm(5.5), Cm(4.5)]
)

heading3(doc, 'Choices (enumeraciones) de cada modelo')
simple_table(doc,
    ['Modelo', 'Campo', 'Opciones'],
    [
        ['Espacio',         'tipo',        'interior_climatizado · exterior_tropical · exterior_cubierto · interior_versatil'],
        ['Espacio',         'estado',      'publicado · borrador'],
        ['Paquete',         'nombre',      'basico · premium · lujo'],
        ['ServicioAdicional','precio_tipo', 'fijo · por_persona'],
        ['FechaOcupada',    'motivo',      'reserva · bloqueo'],
        ['Reserva',         'tipo_evento', 'boda · xv_anos · bautizo · aniversario · empresarial · cumpleanos'],
        ['Reserva',         'estado',      'confirmada · pendiente_liquidacion · completada · cancelada'],
        ['PagoRegistro',    'concepto',    'anticipo · liquidacion · extra'],
        ['TareaPendiente',  'prioridad',   'alta · media · baja'],
    ],
    [Cm(4), Cm(3), Cm(8.5)]
)

body(doc, 'Código PlantUML del diagrama de clases completo (pegar en plantuml.com):')
code_block(doc, '''@startuml Diagrama_Clases
skinparam classAttributeIconSize 0
skinparam classFontSize 11

package "django.contrib.auth" {
  class User {
    +id : Integer
    +username : String
    +email : String
    +is_staff : Boolean
  }
}
package "accounts" {
  class ClienteProfile {
    +telefono : String
    +foto_perfil : ImageField
  }
}
package "venues" {
  class Espacio {
    +nombre; +slug; +tipo; +capacidad_max
    +precio_base; +precio_max; +estado; +orden
    +get_tipo_label(); +tiene_vista_al_mar()
  }
  class EspacioImagen { +imagen; +caption; +orden }
  class EspacioCaracteristica { +descripcion; +orden }
  class Paquete {
    +nombre; +precio_por_persona
    +duracion_horas; +es_mas_popular
  }
  class PaqueteItem { +descripcion; +incluido; +orden }
  class ServicioAdicional {
    +nombre; +precio; +precio_tipo; +activo
    +precio_display()
  }
  class FechaOcupada {
    +fecha; +motivo; +nota
    <<UNIQUE: espacio+fecha>>
  }
}
package "bookings" {
  class Reserva {
    +codigo; +tipo_evento; +fecha; +hora_inicio
    +num_invitados; +estado; +total_evento; +anticipo_pagado
    +saldo_pendiente; +calcular_total(); +get_estado_css()
  }
  class ReservaServicio {
    +cantidad; +subtotal()
    <<UNIQUE: reserva+servicio>>
  }
  class PagoRegistro { +monto; +fecha_pago; +concepto }
  class ActividadReserva { +descripcion; +created_at }
}
package "operations" {
  class TareaPendiente {
    +descripcion; +completada; +prioridad
    +get_color()
  }
}

User "1" --o "0..1" ClienteProfile
User "1" --o "0..*" Reserva : cliente
User "0..1" --o "0..*" Reserva : coordinador
Espacio "1" --o "0..*" EspacioImagen
Espacio "1" --o "0..*" EspacioCaracteristica
Espacio "1" --o "0..*" FechaOcupada
Espacio "1" --o "0..*" Reserva
Paquete "1" --o "1..*" PaqueteItem
Paquete "1" --o "0..*" Reserva
Reserva "1" --o "0..*" ReservaServicio
Reserva "1" --o "1..*" PagoRegistro
Reserva "1" --o "1..*" ActividadReserva
ServicioAdicional "1" --o "0..*" ReservaServicio
User "0..1" --o "0..*" TareaPendiente
@enduml''')

# ── 4.4 ERD ───────────────────────────────────────────────────────────────────
heading2(doc, '4.4 Diagrama de Base de Datos (ERD)')
note_box(doc, 'Diagrama entidad-relación en sintaxis Mermaid. Para renderizar: pegar en https://mermaid.live')

heading3(doc, 'Tablas y relaciones del sistema')
simple_table(doc,
    ['Tabla', 'Columnas principales', 'Foreign Keys'],
    [
        ['auth_user',                    'id, username, email, password, first_name, last_name, is_staff, is_active', '—'],
        ['accounts_clienteprofile',      'id, user_id, telefono, foto_perfil', 'user_id → auth_user'],
        ['venues_espacio',               'id, nombre, slug, tipo, capacidad_max, precio_base, precio_max, estado, orden', '—'],
        ['venues_espacioimagen',         'id, espacio_id, imagen, caption, orden', 'espacio_id → venues_espacio'],
        ['venues_espaciocaracteristica', 'id, espacio_id, descripcion, orden', 'espacio_id → venues_espacio'],
        ['venues_paquete',               'id, nombre, precio_por_persona, duracion_horas, es_mas_popular', '—'],
        ['venues_paqueteitem',           'id, paquete_id, descripcion, incluido, orden', 'paquete_id → venues_paquete'],
        ['venues_servicioadicional',     'id, nombre, precio, precio_tipo, activo, orden', '—'],
        ['venues_fechaocupada',          'id, espacio_id, fecha, motivo, nota | UNIQUE(espacio_id+fecha)', 'espacio_id → venues_espacio'],
        ['bookings_reserva',             'id, codigo(unique), cliente_id, espacio_id, paquete_id, coordinador_id, tipo_evento, fecha, hora_inicio, num_invitados, estado, total_evento, anticipo_pagado, notas_internas', 'cliente_id, espacio_id, paquete_id, coordinador_id → auth_user'],
        ['bookings_reservaservicio',     'id, reserva_id, servicio_id, cantidad | UNIQUE(reserva+servicio)', 'reserva_id, servicio_id'],
        ['bookings_pagoregistro',        'id, reserva_id, registrado_por_id, monto, fecha_pago, concepto', 'reserva_id, registrado_por_id'],
        ['bookings_actividadreserva',    'id, reserva_id, usuario_id, descripcion, created_at', 'reserva_id, usuario_id'],
        ['operations_tareapendiente',    'id, asignado_a_id, descripcion, completada, prioridad, created_at', 'asignado_a_id → auth_user'],
    ],
    [Cm(4.5), Cm(7), Cm(4)]
)

code_block(doc, '''erDiagram
    auth_user {
        int id PK
        varchar username
        varchar email
        bool is_staff
    }
    accounts_clienteprofile {
        int id PK
        int user_id FK
        varchar telefono
    }
    venues_espacio {
        int id PK
        varchar nombre
        varchar slug
        varchar tipo
        int capacidad_max
        decimal precio_base
        varchar estado
    }
    venues_fechaocupada {
        int id PK
        int espacio_id FK
        date fecha
        varchar motivo
    }
    venues_paquete {
        int id PK
        varchar nombre
        decimal precio_por_persona
    }
    venues_servicioadicional {
        int id PK
        varchar nombre
        decimal precio
        varchar precio_tipo
    }
    bookings_reserva {
        int id PK
        varchar codigo
        int cliente_id FK
        int espacio_id FK
        int paquete_id FK
        int coordinador_id FK
        varchar tipo_evento
        date fecha
        int num_invitados
        varchar estado
        decimal total_evento
        decimal anticipo_pagado
    }
    bookings_reservaservicio {
        int id PK
        int reserva_id FK
        int servicio_id FK
        int cantidad
    }
    bookings_pagoregistro {
        int id PK
        int reserva_id FK
        int registrado_por_id FK
        decimal monto
        varchar concepto
        date fecha_pago
    }
    bookings_actividadreserva {
        int id PK
        int reserva_id FK
        int usuario_id FK
        varchar descripcion
        datetime created_at
    }
    operations_tareapendiente {
        int id PK
        int asignado_a_id FK
        varchar descripcion
        bool completada
        varchar prioridad
    }

    auth_user ||--o| accounts_clienteprofile : "tiene perfil"
    auth_user ||--o{ bookings_reserva : "es cliente de"
    auth_user ||--o{ bookings_reserva : "coordina"
    auth_user ||--o{ bookings_pagoregistro : "registra pago"
    auth_user ||--o{ bookings_actividadreserva : "registra"
    auth_user ||--o{ operations_tareapendiente : "asignado a"
    venues_espacio ||--o{ venues_fechaocupada : "fechas bloqueadas"
    venues_espacio ||--o{ bookings_reserva : "es reservado en"
    venues_paquete ||--o{ bookings_reserva : "es contratado"
    bookings_reserva ||--o{ bookings_reservaservicio : "incluye extras"
    bookings_reserva ||--|{ bookings_pagoregistro : "tiene pagos"
    bookings_reserva ||--|{ bookings_actividadreserva : "registra actividad"
    venues_servicioadicional ||--o{ bookings_reservaservicio : "seleccionado en"''', title='Mermaid — Diagrama ERD:')

# ── Pie ───────────────────────────────────────────────────────────────────────
doc.add_paragraph()
hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación generada automáticamente a partir del análisis del código fuente del proyecto.')
run.font.size = Pt(9)
run.font.color.rgb = GRAY
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Renderizadores: PlantUML → plantuml.com  |  Mermaid → mermaid.live')
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY

# ── Guardar ───────────────────────────────────────────────────────────────────
output = r'c:\Users\PC\Complejo-Gestion-Salon-de-Eventos\docs\documentacion_sistema.docx'
doc.save(output)
print(f'Archivo generado: {output}')
